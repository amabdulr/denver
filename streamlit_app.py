"""
Streamlit App - Cisco Documentation Assistant
Main application file for analyzing bug reports and suggesting documentation updates
"""

import streamlit as st
from dotenv import load_dotenv
from bug2 import create_auth, get_bug_summary, get_file_content, get_note_content, get_all_notes, create_note
import xml.etree.ElementTree as ET
import requests
import json
import os

# Import helper functions
from app_functions import run_agent, format_output, apply_prompt_file
from first_draft_tab import render_first_draft_tab, handle_first_draft_button, handle_find_internal_button, handle_explain_sfs_button
from Convert import render_convert_tab, convert_to_xml
from hal_check_tab import render_hal_check_tab, handle_hallucination_check

# Load the .env file
load_dotenv()

# Config file for persistent settings
CONFIG_FILE = "app_config.json"

def load_config():
    """Load application configuration from file"""
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_config(config):
    """Save application configuration to file"""
    try:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f, indent=2)
    except:
        pass

def get_saved_product():
    """Get the saved product name from config"""
    config = load_config()
    return config.get('product_name', 'Cisco 9800')

def save_product_preference(product_name):
    """Save the product name preference"""
    config = load_config()
    config['product_name'] = product_name
    save_config(config)

# Streamlit UI
st.set_page_config(page_title="Cisco Documentation Assistant", page_icon="📚", layout="wide")

# Custom CSS to make tab buttons bigger
st.markdown("""
    <style>
    /* Make tab buttons larger and more prominent */
    button[data-baseweb="tab"] {
        font-size: 24px !important;
        font-weight: 700 !important;
        padding: 16px 32px !important;
        height: auto !important;
        min-height: 60px !important;
    }
    
    /* Style the tab list container */
    div[data-baseweb="tab-list"] {
        gap: 12px !important;
    }
    
    /* Add spacing between tab button text and icons */
    button[data-baseweb="tab"] > div {
        gap: 10px !important;
    }
    
    /* Make the tab panel border more visible */
    div[data-baseweb="tab-panel"] {
        padding-top: 20px !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("🔧 Writer Workflow")
st.markdown("*Analyze bugs and suggest documentation updates using AI-powered RAG*")

# Initialize vector store on first load (auto-detects SQLite version for persistence)
if 'vector_store_initialized' not in st.session_state:
    with st.spinner("🔄 Loading knowledge base..."):
        try:
            from vector_store_manager import initialize_vector_store, is_initialized, get_persistence_mode
            if not is_initialized():
                initialize_vector_store()
            
            mode = get_persistence_mode()
            st.session_state.vector_store_initialized = True
            st.session_state.persistence_mode = mode
            
            if mode == 'persistent':
                st.success("✅ Knowledge base loaded (persistent mode)", icon="✅")
            else:
                st.success("✅ Knowledge base loaded (in-memory mode)", icon="✅")
                st.info("ℹ️ Running in-memory mode due to SQLite version. Data will not persist between restarts.")
        except Exception as e:
            st.error(f"❌ Error loading knowledge base: {e}")
            st.info("💡 Make sure knowledge_docs/ directory exists and contains files")
            st.session_state.vector_store_initialized = False

# Create tabs at the top for different workflows
tab1, tab2, tab3 = st.tabs(["🔍 Analysis & Summary", "✍️ First Draft", "🔍 Hal-Check"])

st.divider()

# Initialize session state for conversation history
if 'conversation_history' not in st.session_state:
    st.session_state.conversation_history = []
if 'initial_analysis_done' not in st.session_state:
    st.session_state.initial_analysis_done = False
if 'current_rca_content' not in st.session_state:
    st.session_state.current_rca_content = ""
if 'uploaded_file_content' not in st.session_state:
    st.session_state.uploaded_file_content = ""
if 'context_window_size' not in st.session_state:
    st.session_state.context_window_size = 3  # Default: last 3 exchanges

def estimate_tokens(text):
    """Rough token estimation: ~4 chars per token"""
    return len(text) // 4

def get_relevant_context(conversation_history, window_size=3, max_tokens=4000):
    """
    Get relevant context from conversation history with token limits
    
    Args:
        conversation_history: List of Q&A exchanges
        window_size: Number of recent exchanges to include
        max_tokens: Maximum tokens for context
    
    Returns:
        Formatted context string
    """
    if not conversation_history:
        return ""
    
    # Get last N exchanges
    recent_exchanges = conversation_history[-window_size:] if len(conversation_history) > window_size else conversation_history
    
    # Build context with token awareness
    context_parts = []
    total_tokens = 0
    
    for idx, exchange in enumerate(reversed(recent_exchanges), 1):
        # Format exchange
        exchange_text = f"Previous Q{len(recent_exchanges)-idx+1}: {exchange['question']}\nPrevious A{len(recent_exchanges)-idx+1}: {exchange['answer'][:1000]}...\n"  # Limit answer preview
        
        tokens = estimate_tokens(exchange_text)
        if total_tokens + tokens > max_tokens:
            break
        
        context_parts.insert(0, exchange_text)
        total_tokens += tokens
    
    if context_parts:
        return "RECENT CONVERSATION CONTEXT:\n" + "\n".join(context_parts)
    return ""

def build_followup_prompt(followup_question, context, use_rag):
    """
    Build a well-structured follow-up prompt
    
    Args:
        followup_question: The user's follow-up question
        context: Previous conversation context
        use_rag: Whether RAG search is enabled
    
    Returns:
        Formatted prompt string
    """
    if use_rag:
        # For RAG: Focus on new search with context awareness
        prompt = f"""You are answering a follow-up question in an ongoing conversation about Cisco documentation.

{context}

CURRENT FOLLOW-UP QUESTION: {followup_question}

INSTRUCTIONS:
- You have context from previous exchanges above
- Search the documentation database for information relevant to this follow-up
- Reference previous answers if relevant (e.g., "As mentioned earlier...")
- If this question asks for clarification/expansion of a previous answer, identify what to expand
- Provide a direct, focused answer to the follow-up question
- Do not repeat information already provided unless specifically asked

Your answer:"""
    else:
        # For direct LLM: Pure conversational follow-up
        prompt = f"""You are continuing a conversation about Cisco documentation and bug analysis.

{context}

USER FOLLOW-UP: {followup_question}

INSTRUCTIONS:
- Answer based on the conversation context above and your general knowledge
- Reference previous exchanges when relevant
- If asking for clarification, expand on the specific point mentioned
- Keep answers concise and focused on what was asked
- Use natural conversational language

Your answer:"""
    
    return prompt

# Shared inputs (outside tabs) - These will be rendered inside each tab
def render_shared_inputs(tab_prefix="", show_header=True):
    """Render inputs that are shared across all tabs"""
    if show_header:
        st.subheader("📝 Input Configuration")
    
    # CDETS Bug Fetcher Section
    st.markdown("#### 🐛 Fetch Bug from CDETS")
    
    # Add checkbox for extracting all notes
    extract_all_notes = st.checkbox(
        "📋 Extract all notes (default: Behavior-changed + Release-note)",
        value=False,
        help="Check this to extract all notes from the bug. By default, only 'Behavior-changed' and 'Release-note' notes are extracted along with the bug summary.",
        key=f"{tab_prefix}_extract_all_notes"
    )
    
    bug_col1, bug_col2 = st.columns([3, 1])
    with bug_col1:
        bug_number_input = st.text_input(
            "Bug Number(s)",
            placeholder="e.g., CSCwp05354 or CSCwp05354, CSCwp12345, CSCwp67890",
            help="Enter one or more CDETS bug numbers (comma-separated) to fetch bug details",
            key=f"{tab_prefix}_bug_number"
        )
    with bug_col2:
        st.write("")  # Spacer
        fetch_bug_button = st.button("🔍 Fetch Bug(s)", use_container_width=True, key=f"{tab_prefix}_fetch_bug")
    
    # Show success message after rerun
    if 'bug_fetched' in st.session_state and st.session_state.bug_fetched:
        st.success(f"✅ Successfully fetched bug data!")
        st.session_state.bug_fetched = False  # Clear the flag
    
    if fetch_bug_button and bug_number_input:
        # Parse multiple bug numbers (comma-separated)
        bug_numbers = [bug.strip() for bug in bug_number_input.split(',') if bug.strip()]
        
        with st.spinner(f"Fetching {len(bug_numbers)} bug(s) from CDETS..."):
            try:
                auth = create_auth()
                ns = {'cdets': 'cdetsng', 'ns2': 'http://www.w3.org/1999/xlink'}
                
                # Combined content for all bugs
                all_bugs_content = ""
                all_notes_summary = {}  # Track notes per bug for display
                
                for bug_idx, bug_number in enumerate(bug_numbers, 1):
                    # Add separator between bugs
                    if bug_idx > 1:
                        all_bugs_content += "\n\n" + "="*80 + "\n\n"
                    
                    # Get bug summary
                    summary_response = get_bug_summary(bug_number, auth)
                    summary_root = ET.fromstring(summary_response.content)
                    
                    # Build bug content
                    bug_content = f"# Bug {bug_number} - Complete Report\n\n"
                    bug_content += "## Bug Summary\n\n"
                    
                    # Extract bug fields
                    defect = summary_root.find('.//cdets:Defect', ns)
                    if defect:
                        for field in defect.findall('.//cdets:Field', ns):
                            field_name = field.get('name')
                            field_value = field.text if field.text else 'N/A'
                            
                            if field_name in ['Headline', 'Status', 'Severity', 'Priority', 'Product', 
                                             'Component', 'Version', 'Description', 'FoundIn', 'FixedIn']:
                                bug_content += f"**{field_name}:** {field_value}\n\n"
                    
                    # Get notes
                    bug_content += "\n## Notes\n\n"
                    
                    note_titles = []  # Collect note titles for this bug
                    
                    if extract_all_notes:
                        # Extract all notes
                        try:
                            all_note_titles = get_all_notes(bug_number, auth)
                            note_titles = all_note_titles
                            
                            for i, note_title in enumerate(all_note_titles, 1):
                                try:
                                    note_response = get_note_content(bug_number, note_title, auth)
                                    bug_content += f"### {i}. {note_title}\n\n"
                                    bug_content += f"**Content:**\n{note_response.text}\n\n"
                                except Exception as e:
                                    bug_content += f"*Error fetching note '{note_title}': {str(e)}*\n\n"
                        except Exception as e:
                            bug_content += f"*Error fetching notes list: {str(e)}*\n\n"
                    else:
                        # Extract Behavior-changed and Release-note by default
                        default_notes = ["Behavior-changed", "Release-note"]
                        note_titles = []
                        
                        for note_title in default_notes:
                            try:
                                note_response = get_note_content(bug_number, note_title, auth)
                                bug_content += f"### {note_title}\n\n"
                                bug_content += f"**Content:**\n{note_response.text}\n\n"
                                note_titles.append(note_title)
                            except Exception as e:
                                # If note doesn't exist, skip silently
                                bug_content += f"*Note '{note_title}' not found*\n\n"
                    
                    # Store note titles for this bug
                    all_notes_summary[bug_number] = note_titles
                    
                    # Append this bug's content to the combined content
                    all_bugs_content += bug_content
                
                # Store in session state
                st.session_state.uploaded_file_content = all_bugs_content
                st.session_state.fetched_notes_summary = all_notes_summary  # Store notes summary
                st.session_state.bug_fetched = True  # Flag to show success message
                # Set the text area value directly
                st.session_state.rca_text_area = all_bugs_content
                st.rerun()
                
            except Exception as e:
                st.error(f"❌ Error fetching bug(s): {str(e)}")
                with st.expander("🐛 Error Details"):
                    st.exception(e)
    
    st.markdown("---")
    
    # Product name input with persistence
    product_options = ["Cisco SD-WAN", "Cisco 9800", "ASR 9000", "Cisco 8000", "cisco_generic"]
    saved_product = get_saved_product()
    
    # Find index of saved product, default to 1 if not found
    try:
        default_index = product_options.index(saved_product)
    except ValueError:
        default_index = 1
    
    product_name = st.selectbox(
        "Product Name",
        options=product_options,
        index=default_index,
        help="Select the Cisco product (selection is remembered)",
        key=f"{tab_prefix}_product_name",
        on_change=lambda: save_product_preference(st.session_state.get(f"{tab_prefix}_product_name"))
    )
    
    return product_name

# Tab 1 - Analysis & Summary
with tab1:
    col1, col2 = st.columns(2)
    
    with col1:
        product_name = render_shared_inputs("tab1")
        
        # Load default prompt from BugAnalyze.md
        try:
            with open("BugAnalyze.md", "r") as f:
                default_prompt = f.read()
        except FileNotFoundError:
            default_prompt = "Analyze the Bug/RCA content"
        
        # Question input for Analysis (only in Tab 1)
        question = st.text_area(
            "Question/Task",
            value=default_prompt,
            key="analysis_question",
            height=200
        )
        
        st.markdown("---")
        
        # Display notes summary if available (persistent display after fetch)
        if 'fetched_notes_summary' in st.session_state and st.session_state.fetched_notes_summary:
            all_notes_summary = st.session_state.fetched_notes_summary
            total_notes = sum(len(notes) for notes in all_notes_summary.values())
            
            with st.expander(f"📋 Extracted Notes Summary ({total_notes} total notes)", expanded=False):
                for bug_num, notes in all_notes_summary.items():
                    st.markdown(f"**Bug {bug_num}** ({len(notes)} notes):")
                    for idx, note_title in enumerate(notes, 1):
                        st.markdown(f"  {idx}. {note_title}")
                    if bug_num != list(all_notes_summary.keys())[-1]:  # Not the last bug
                        st.markdown("---")
        
        # RCA content input
        rca_content = st.text_area(
            "Bug Report / RCA Content",
            height=300,
            placeholder="Paste your bug report, root cause analysis, or case notes here...",
            help="Paste the content of your bug report or RCA",
            key="rca_text_area"
        )
    
    with col2:
        st.subheader("📊 Output")
        
        # Add "Post Analysis to Bug" button above output
        post_analysis_button = st.button("📤 Post Analysis to Bug", type="secondary", use_container_width=True, key="post_analysis_to_bug")
        
        output_container = st.container()
        
        # Display follow-up answer if available
        if 'last_followup_answer' in st.session_state and st.session_state.last_followup_answer:
            with output_container:
                st.markdown("## 💬 Follow-up Answer")
                st.markdown(st.session_state.last_followup_answer)
                
                # Show raw output in expander
                if 'last_followup_raw' in st.session_state:
                    with st.expander("🔍 View Raw Response"):
                        st.json(st.session_state.last_followup_raw)
    
    st.divider()
    st.markdown("### 🔍 Analysis & Summary Tools")
    
    # Analysis and Summary buttons
    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 4])

    with col_btn1:
        summarize_button = st.button("📝 Summarize", type="primary", use_container_width=True)

    with col_btn2:
        analyze_button = st.button("🚀 Analyze", use_container_width=True)

    with col_btn3:
        clear_button = st.button("🗑️ Clear", use_container_width=True)

# Tab 2 - First Draft
with tab2:
    # Add custom styling and header
    st.markdown("""
        <style>
        .big-title {
            font-size: 28px;
            font-weight: 700;
            color: #1f77b4;
            margin-bottom: 10px;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<p class="big-title">✍️ First Draft Generation</p>', unsafe_allow_html=True)
    st.markdown("Transform your SFS content into clean, customer-facing documentation")
    
    # Initialize session state for button state tracking
    if 'find_internal_clicked' not in st.session_state:
        st.session_state.find_internal_clicked = False
    
    # Workflow steps indicator
    if not st.session_state.find_internal_clicked:
        st.info("📍 **Step 1 of 2:** First, identify internal information in your SFS document")
    else:
        st.success("✅ Step 1 complete | 📍 **Step 2 of 2:** Ready to generate your first draft")
    
    st.markdown("---")
    
    col1_draft, col2_draft = st.columns(2)
    
    with col1_draft:
        # File uploader for Bug/RCA content (only in First Draft tab)
        st.markdown("#### 📎 Upload Document(s)")
        uploaded_files = st.file_uploader(
            "Upload Bug/RCA Document(s)",
            type=["txt", "md", "log", "doc", "docx", "pdf"],
            help="Upload one or more files to auto-populate the SFS field below",
            key="first_draft_uploader",
            label_visibility="collapsed",
            accept_multiple_files=True
        )
        
        # Button to populate from uploaded files
        if uploaded_files:
            if st.button("📥 Paste File Content to SFS Field", type="secondary", use_container_width=True, key="paste_file_draft"):
                try:
                    import io
                    all_content = []
                    
                    for uploaded_file in uploaded_files:
                        # Read file content based on type
                        file_extension = uploaded_file.name.split('.')[-1].lower()
                        
                        # Add file separator
                        all_content.append(f"\n{'='*80}\n")
                        all_content.append(f"FILE: {uploaded_file.name}\n")
                        all_content.append(f"{'='*80}\n\n")
                        
                        if file_extension in ['txt', 'md', 'log']:
                            # Read text files directly
                            content = uploaded_file.read().decode('utf-8')
                            all_content.append(content)
                        elif file_extension in ['doc', 'docx']:
                            # Try to import and use python-docx
                            try:
                                from docx import Document as DocxDocument
                                # Read the docx file
                                doc = DocxDocument(io.BytesIO(uploaded_file.read()))
                                # Extract text from all paragraphs
                                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                                all_content.append(content)
                            except ImportError:
                                st.warning(f"⚠️ Word document support requires 'python-docx' for {uploaded_file.name}")
                                st.info("Please run: pip install python-docx")
                                continue
                        elif file_extension == 'pdf':
                            st.warning(f"⚠️ PDF file {uploaded_file.name} requires additional libraries. Please copy and paste the content or save as .txt")
                            continue
                        else:
                            st.warning(f"⚠️ Unsupported file type: {file_extension} for {uploaded_file.name}")
                            continue
                    
                    # Combine all content
                    combined_content = ''.join(all_content)
                    st.session_state.uploaded_file_content = combined_content
                    st.session_state.extracted_text_area = combined_content  # Update the keyed widget state
                    st.success(f"✅ Pasted {len(combined_content)} characters from {len(uploaded_files)} file(s)")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Error reading files: {str(e)}")
        
        st.markdown("---")
        
        # Product name input only (no fetch bug functionality in Tab 2) with persistence
        product_options = ["Cisco SD-WAN", "Cisco 9800", "ASR 9000", "Cisco 8000", "cisco_generic"]
        saved_product = get_saved_product()
        
        # Find index of saved product, default to 1 if not found
        try:
            default_index = product_options.index(saved_product)
        except ValueError:
            default_index = 1
        
        product_name_draft = st.selectbox(
            "Product Name",
            options=product_options,
            index=default_index,
            help="Select the Cisco product (selection is remembered)",
            key="tab2_product_name",
            on_change=lambda: save_product_preference(st.session_state.get("tab2_product_name"))
        )
        
        st.markdown("---")
        
        # SFS content input (for Tab 2 - First Draft)
        # Use current_extracted_text if it exists (after cleaning), otherwise use the widget's value
        default_value = st.session_state.get('current_extracted_text', '')
        extracted_text = st.text_area(
            "SFS",
            value=default_value,
            height=300,
            placeholder="Paste your SFS content here...",
            help="Paste the content of your SFS",
            key="extracted_text_area"
        )
    
    with col2_draft:
        st.subheader("📊 Output")
        output_container_draft = st.container()
        
        # Display follow-up answer if available
        if 'last_followup_answer' in st.session_state and st.session_state.last_followup_answer:
            with output_container_draft:
                st.markdown("## 💬 Follow-up Answer")
                st.markdown(st.session_state.last_followup_answer)
                
                # Show raw output in expander
                if 'last_followup_raw' in st.session_state:
                    with st.expander("🔍 View Raw Response"):
                        st.json(st.session_state.last_followup_raw)
    
    st.divider()
    
    # Render the First Draft tab from the imported module
    first_draft_button, clear_button_draft, find_internal_button, explain_sfs_button = render_first_draft_tab(output_container_draft)

# Handle button clicks
if analyze_button:
    if not rca_content.strip():
        st.error("⚠️ Please provide RCA content to analyze.")
    elif not product_name.strip():
        st.error("⚠️ Please provide a product name.")
    elif not question.strip():
        st.error("⚠️ Please provide a question or task description.")
    else:
        with st.spinner("🔍 Analyzing documentation and generating recommendations..."):
            try:
                # Store in session state
                st.session_state.product_name = product_name
                st.session_state.current_rca_content = rca_content
                st.session_state.initial_analysis_done = True
                
                # Run the agent
                result = run_agent(product_name, question, rca_content)
                
                # Add to conversation history
                st.session_state.conversation_history.append({
                    "question": question,
                    "answer": result['output'] if 'output' in result else str(result)
                })
                
                # Display formatted output
                with output_container:
                    formatted_output = format_output(result)
                    st.markdown(formatted_output)
                    
                    # Show raw output in expander
                    with st.expander("🔍 View Raw Response"):
                        st.json(result)
                        
                st.success("✅ Analysis complete!")
                
            except Exception as e:
                st.error(f"❌ Error during analysis: {str(e)}")
                with st.expander("🐛 Error Details"):
                    st.exception(e)

if post_analysis_button:
    # Get the first bug number from the input
    bug_number_input = st.session_state.get('tab1_bug_number', '')
    
    if not bug_number_input:
        st.error("⚠️ Please enter a bug number first.")
    else:
        # Get the first bug number (in case multiple were entered)
        first_bug = bug_number_input.split(',')[0].strip()
        
        # Get the output content from conversation history
        if not st.session_state.conversation_history:
            st.error("⚠️ No analysis output found. Please run an analysis first.")
        else:
            # Get the last answer as the note body
            last_answer = st.session_state.conversation_history[-1]['answer']
            
            with st.spinner(f"📤 Posting analysis to bug {first_bug}..."):
                try:
                    auth = create_auth()
                    response = create_note(
                        bug_number=first_bug,
                        note_title="AI-Analysis",
                        note_content=last_answer,
                        note_type="Other",
                        auth=auth
                    )
                    st.success(f"✅ Successfully posted analysis to bug {first_bug}!")
                    st.info(f"Response status: {response.status_code}")
                    
                except Exception as e:
                    st.error(f"❌ Error posting analysis to bug: {str(e)}")
                    with st.expander("🐛 Error Details"):
                        st.exception(e)

if summarize_button:
    if not rca_content.strip():
        st.error("⚠️ Please provide RCA content to summarize.")
    else:
        with st.spinner("📝 Generating summary..."):
            try:
                # Store in session state
                st.session_state.product_name = product_name
                st.session_state.current_rca_content = rca_content
                st.session_state.initial_analysis_done = True
                
                # Use the apply_prompt_file function with summarize.md
                summary = apply_prompt_file("summarize.md", rca_content, product_name)
                
                # Add to conversation history
                st.session_state.conversation_history.append({
                    "question": "Summarize the bug/RCA content",
                    "answer": summary
                })
                
                with output_container:
                    st.markdown("## 📝 Summary")
                    st.markdown(summary)
                
                st.success("✅ Summary generated!")
                
            except Exception as e:
                st.error(f"❌ Error generating summary: {str(e)}")
                with st.expander("🐛 Error Details"):
                    st.exception(e)

if first_draft_button:
    handle_first_draft_button(extracted_text, product_name_draft, output_container_draft)

if find_internal_button:
    handle_find_internal_button(extracted_text, product_name_draft, output_container_draft)

if explain_sfs_button:
    handle_explain_sfs_button(extracted_text, product_name_draft, output_container_draft)

if clear_button or clear_button_draft:
    st.session_state.conversation_history = []
    st.session_state.initial_analysis_done = False
    st.session_state.current_rca_content = ""
    st.session_state.current_extracted_text = ""
    st.session_state.uploaded_file_content = ""
    # Clear follow-up answer display
    if 'last_followup_answer' in st.session_state:
        del st.session_state.last_followup_answer
    if 'last_followup_raw' in st.session_state:
        del st.session_state.last_followup_raw
    # Reset button state for Find Internal Information
    if 'find_internal_clicked' in st.session_state:
        st.session_state.find_internal_clicked = False
    # Clear text area widgets by deleting their session state keys
    if 'rca_text_area' in st.session_state:
        del st.session_state.rca_text_area
    if 'extracted_text_area' in st.session_state:
        del st.session_state.extracted_text_area
    # Also delete current_extracted_text to ensure text area clears
    if 'current_extracted_text' in st.session_state:
        del st.session_state.current_extracted_text
    st.rerun()

# Initialize session state for tracking which tab to show follow-up output
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = 'tab1'

# Follow-up question section
if st.session_state.initial_analysis_done and st.session_state.conversation_history:
    st.divider()
    st.subheader("💬 Follow-up Questions")
    
    # Settings for context window
    with st.expander("⚙️ Follow-up Settings", expanded=False):
        context_window = st.slider(
            "Context Window (number of recent exchanges to include)",
            min_value=1,
            max_value=10,
            value=st.session_state.context_window_size,
            help="Controls how many recent Q&A exchanges are included in follow-up context. Lower = faster, higher = more context."
        )
        st.session_state.context_window_size = context_window
        
        # Show token estimate
        if st.session_state.conversation_history:
            recent_context = get_relevant_context(st.session_state.conversation_history, context_window)
            est_tokens = estimate_tokens(recent_context)
            st.info(f"📊 Current context: ~{est_tokens} tokens from last {min(context_window, len(st.session_state.conversation_history))} exchanges")
    
    # Display conversation history
    with st.expander(f"📜 View Conversation History ({len(st.session_state.conversation_history)} exchanges)", expanded=False):
        for idx, exchange in enumerate(st.session_state.conversation_history, 1):
            st.markdown(f"**Q{idx}:** {exchange['question']}")
            st.markdown(f"**A{idx}:** {exchange['answer']}")
            st.markdown("---")
    
    # Follow-up question input
    followup_col1, followup_col2 = st.columns([4, 1])
    
    with followup_col1:
        followup_question = st.text_area(
            "Ask a follow-up question",
            placeholder="e.g., Can you provide more details about the configuration section?",
            height=100,
            key="followup_input"
        )
        
        # Checkbox for RAG search
        use_rag_search = st.checkbox(
            "🔍 Search documentation for this follow-up",
            value=False,
            help="Enable to search documentation database (like Analyze). Disable for conversational answers only."
        )
    
    with followup_col2:
        st.write("")  # Spacer
        st.write("")  # Spacer
        submit_followup = st.button("📤 Ask", type="secondary", use_container_width=True)
    
    if submit_followup:
        if not followup_question.strip():
            st.warning("⚠️ Please enter a follow-up question.")
        else:
            with st.spinner("🔍 Processing your follow-up question..."):
                try:
                    # Get relevant context with token management
                    context = get_relevant_context(
                        st.session_state.conversation_history, 
                        st.session_state.context_window_size,
                        max_tokens=4000
                    )
                    
                    # Build structured prompt
                    full_followup = build_followup_prompt(followup_question, context, use_rag_search)
                    
                    # Determine which content to use based on which tab is active
                    content_to_use = st.session_state.get('current_extracted_text') or st.session_state.get('current_rca_content', '')
                    
                    # Choose between RAG search or direct LLM based on checkbox
                    if use_rag_search:
                        # Use RAG search with improved prompt
                        result = run_agent(
                            st.session_state.product_name, 
                            full_followup, 
                            content_to_use
                        )
                        followup_answer = result['output'] if 'output' in result else str(result)
                    else:
                        # Direct LLM call with structured prompt (no RAG)
                        from utils import get_llm
                        llm = get_llm()
                        # Add original content if needed
                        if content_to_use:
                            full_followup += f"\n\nORIGINAL BUG/RCA CONTENT (for reference):\n{content_to_use[:2000]}..."
                        result = llm.invoke(full_followup)
                        followup_answer = result.content if hasattr(result, 'content') else str(result)
                    
                    # Add to conversation history
                    st.session_state.conversation_history.append({
                        "question": followup_question,
                        "answer": followup_answer
                    })
                    
                    # Store the follow-up answer to display in output container
                    st.session_state.last_followup_answer = followup_answer
                    st.session_state.last_followup_raw = result if use_rag_search else {"output": followup_answer}
                    
                    st.success("✅ Follow-up answered! Check the Output section above.")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Error processing follow-up: {str(e)}")
                    with st.expander("🐛 Error Details"):
                        st.exception(e)

# Tab 3 - Hallucination Check
with tab3:
    render_hal_check_tab()

# Tab 4 - Convert (Hidden temporarily)
# with tab3:
#     st.markdown("### 🔄 Convert Content to XML Format")
#     st.markdown("Convert your raw content to various XML information types.")
#     
#     # Handle convert button BEFORE rendering to ensure state is updated
#     # Get the raw content and conversion type from session state for button handling
#     if 'convert_triggered' in st.session_state and st.session_state.convert_triggered:
#         raw_content_to_convert = st.session_state.get('pending_raw_content', '')
#         conversion_type_to_use = st.session_state.get('pending_conversion_type', 'ct-concept')
#         
#         if raw_content_to_convert.strip():
#             with st.spinner(f"🔄 Converting to {conversion_type_to_use} XML format..."):
#                 try:
#                     # Perform the conversion
#                     xml_output, error = convert_to_xml(
#                         raw_content_to_convert, 
#                         conversion_type_to_use, 
#                         "ConvertToXML.md"
#                     )
#                     
#                     if error:
#                         st.error(f"❌ {error}")
#                     else:
#                         # Store the output in session state
#                         st.session_state.xml_output = xml_output
#                         st.success(f"✅ Successfully converted to {conversion_type_to_use} XML!")
#                         
#                 except Exception as e:
#                     st.error(f"❌ Error during conversion: {str(e)}")
#                     with st.expander("🐛 Error Details"):
#                         st.exception(e)
#         
#         # Reset the trigger
#         st.session_state.convert_triggered = False
#     
#     # Render the Convert tab
#     convert_button, raw_content, conversion_type = render_convert_tab()
#     
#     # Handle convert button click
#     if convert_button:
#         if not raw_content.strip():
#             st.error("⚠️ Please provide raw content to convert.")
#         else:
#             # Store the content for next rerun
#             st.session_state.pending_raw_content = raw_content
#             st.session_state.pending_conversion_type = conversion_type
#             st.session_state.convert_triggered = True
#             st.rerun()


