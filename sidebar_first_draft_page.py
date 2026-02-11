"""
First Draft Sidebar Page Component
Handles the First Draft workflow for sidebar navigation app
"""

import streamlit as st
import io
from app_functions import apply_prompt_file, format_output


def render_first_draft_page():
    """Render the First Draft page with full functionality for sidebar navigation"""
    st.header("✍️ First Draft Generation")
    st.markdown("Transform your SFS content into clean, customer-facing documentation")
    
    # Step-by-step instructions
    with st.expander("📖 How to Use This Page", expanded=False):
        st.markdown("""
        #### Two-Step Workflow for First Draft Generation
        
        **Step 1: Identify Internal Information**
        1. **Upload Document(s)**: Upload your SFS document(s) and Meeting Transcripts. (.txt, .md, .doc, .docx, .pdf)
        2. **Paste Content**: Click "📥 Paste File Content to SFS Field" to populate the text area
        3. **Select Product**: Choose your Cisco product from the dropdown
        4. **Run Analysis**: Click "🔍 Find Internal Information" to identify internal-only content
        5. **Review Results**: Check the identified internal information in the right column
                    
        
        **Step 2: Generate First Draft**
        1. **Review SFS Content**: Make sure your SFS content is complete in the left column
        2. **Click Generate**: Click "✨ Generate First Draft" to create customer-facing documentation
        3. **Review Output**: Check the generated draft in the right column
        
        **Tips**:
        - The workflow is sequential - complete Step 1 before Step 2
        - Internal information findings help guide the first draft generation
        - You can edit the SFS content manually before generating the draft
        - Multiple files can be uploaded and combined
        - The Clear button resets the workflow to start fresh.            
                    
        """)
    
    # Initialize session state for button state tracking
    if 'find_internal_clicked' not in st.session_state:
        st.session_state.find_internal_clicked = False
    
    # Workflow steps indicator
    if not st.session_state.find_internal_clicked:
        st.info("📍 **Step 1 of 2:** First, identify internal information in your SFS document")
    else:
        st.success("✅ Step 1 complete | 📍 **Step 2 of 2:** Ready to generate your first draft")
    
    st.markdown("---")
    
    col1_draft, col2_draft = st.columns(2)
    
    with col1_draft:
        # File uploader for Bug/RCA content
        st.markdown("#### 📎 Upload Document(s)")
        uploaded_files = st.file_uploader(
            "Upload Bug/RCA Document(s)",
            type=["txt", "md", "log", "doc", "docx", "pdf"],
            help="Upload one or more files to auto-populate the SFS field below",
            key="sidebar_first_draft_uploader",
            label_visibility="collapsed",
            accept_multiple_files=True
        )
        
        # Button to populate from uploaded files
        if uploaded_files:
            if st.button("📥 Paste File Content to SFS Field", type="secondary", use_container_width=True, key="sidebar_paste_file_draft"):
                try:
                    all_content = []
                    
                    for uploaded_file in uploaded_files:
                        # Read file content based on type
                        file_extension = uploaded_file.name.split('.')[-1].lower()
                        
                        # Add file separator
                        all_content.append(f"\n{'='*80}\n")
                        all_content.append(f"FILE: {uploaded_file.name}\n")
                        all_content.append(f"{'='*80}\n\n")
                        
                        if file_extension in ['txt', 'md', 'log']:
                            # Read text files directly
                            content = uploaded_file.read().decode('utf-8')
                            all_content.append(content)
                        elif file_extension in ['doc', 'docx']:
                            # Try to import and use python-docx
                            try:
                                from docx import Document as DocxDocument
                                # Read the docx file
                                doc = DocxDocument(io.BytesIO(uploaded_file.read()))
                                # Extract text from all paragraphs
                                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                                all_content.append(content)
                            except ImportError:
                                st.warning(f"⚠️ Word document support requires 'python-docx' for {uploaded_file.name}")
                                st.info("Please run: pip install python-docx")
                                continue
                        elif file_extension == 'pdf':
                            st.warning(f"⚠️ PDF file {uploaded_file.name} requires additional libraries. Please copy and paste the content or save as .txt")
                            continue
                        else:
                            st.warning(f"⚠️ Unsupported file type: {file_extension} for {uploaded_file.name}")
                            continue
                    
                    # Combine all content
                    combined_content = ''.join(all_content)
                    st.session_state.uploaded_file_content = combined_content
                    st.session_state.sidebar_extracted_text_area = combined_content  # Update the keyed widget state
                    st.success(f"✅ Pasted {len(combined_content)} characters from {len(uploaded_files)} file(s)")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Error reading files: {str(e)}")
        
        st.markdown("---")
        
        # Product name input only (no fetch bug functionality in First Draft)
        from sidebar_app import get_saved_product, save_product_preference
        
        product_options = ["Cisco SD-WAN", "Cisco 9800", "ASR 9000", "Cisco 8000", "cisco_generic"]
        saved_product = get_saved_product()
        
        # Find index of saved product, default to 1 if not found
        try:
            default_index = product_options.index(saved_product)
        except ValueError:
            default_index = 1
        
        product_name_draft = st.selectbox(
            "Product Name",
            options=product_options,
            index=default_index,
            help="Select the Cisco product (selection is remembered)",
            key="sidebar_first_draft_product_name",
            on_change=lambda: save_product_preference(st.session_state.get("sidebar_first_draft_product_name"))
        )
        
        st.markdown("---")
        
        # SFS content input
        # Use current_extracted_text if it exists (after cleaning), otherwise use the widget's value
        default_value = st.session_state.get('current_extracted_text', '')
        extracted_text = st.text_area(
            "SFS",
            value=default_value,
            height=300,
            placeholder="Paste your SFS content here...",
            help="Paste the content of your SFS",
            key="sidebar_extracted_text_area"
        )
    
    with col2_draft:
        st.subheader("📊 Output")
        output_container_draft = st.container()
        
        # Display follow-up answer if available
        if 'last_followup_answer' in st.session_state and st.session_state.last_followup_answer:
            with output_container_draft:
                st.markdown("## 💬 Follow-up Answer")
                st.markdown(st.session_state.last_followup_answer)
                
                # Show raw output in expander
                if 'last_followup_raw' in st.session_state:
                    with st.expander("🔍 View Raw Response"):
                        st.json(st.session_state.last_followup_raw)
    
    # Note: Follow-up section will be rendered AFTER button handlers execute
    # This ensures outputs appear before the follow-up interface
    
    st.divider()
    
    # Add custom styling for section headers
    st.markdown("""
        <style>
        .section-header {
            font-size: 18px;
            font-weight: 600;
            color: #2c3e50;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<p class="section-header">🎯 Actions</p>', unsafe_allow_html=True)
    
    # First Draft buttons with dynamic styling based on state
    col_draft1, col_draft2, col_draft3, col_draft4 = st.columns([1.2, 1.2, 1.2, 1])
    
    # Determine button types based on state
    find_button_type = "secondary" if st.session_state.find_internal_clicked else "primary"
    draft_button_type = "primary" if st.session_state.find_internal_clicked else "secondary"
    
    with col_draft1:
        explain_sfs_button = st.button(
            "💡 Explain the SFS",
            type="primary",
            use_container_width=True,
            help="Get an explanation of the SFS content",
            key="sidebar_explain_sfs"
        )
    
    with col_draft2:
        find_internal_button = st.button(
            "🔍 Find Internal Information", 
            type=find_button_type, 
            use_container_width=True,
            help="Analyze SFS content to identify internal/confidential information",
            key="sidebar_find_internal"
        )
    
    with col_draft3:
        # Generate First Draft button
        first_draft_button = st.button(
            "✍️ Generate First Draft", 
            type=draft_button_type, 
            use_container_width=True,
            help="Create a customer-facing draft",
            key="sidebar_first_draft"
        )
    
    with col_draft4:
        clear_button_draft = st.button(
            "🗑️ Clear", 
            use_container_width=True, 
            key="sidebar_clear_draft",
            help="Reset and start over"
        )
    
    # ===== HANDLE BUTTON CLICKS =====
    
    # Handle Explain SFS button
    if explain_sfs_button:
        if not extracted_text.strip():
            st.error("⚠️ Please provide SFS content to explain.")
        elif not product_name_draft.strip():
            st.error("⚠️ Please select a product name.")
        else:
            with st.spinner("💡 Explaining SFS content..."):
                try:
                    # Store in session state
                    st.session_state.product_name = product_name_draft
                    st.session_state.current_extracted_text = extracted_text
                    st.session_state.initial_analysis_done = True
                    
                    # Use the apply_prompt_file function with SFSExplainer.md
                    explanation = apply_prompt_file("SFSExplainer.md", extracted_text, product_name_draft)
                    
                    # Add to conversation history
                    st.session_state.conversation_history.append({
                        "question": "Explain the SFS",
                        "answer": explanation
                    })
                    
                    with output_container_draft:
                        st.markdown("## 💡 SFS Explanation")
                        st.markdown(explanation)

                except Exception as e:
                    st.error(f"❌ Error explaining SFS: {str(e)}")
                    with st.expander("🐛 Error Details"):
                        st.exception(e)
    
    # Handle Find Internal Information button
    if find_internal_button:
        # Clear previous conversation context when starting Find Internal workflow
        if st.session_state.conversation_history:
            # Check if last conversation was Explain SFS
            if st.session_state.conversation_history[-1].get('question') == 'Explain the SFS':
                st.session_state.conversation_history = []
        
        if not extracted_text.strip():
            st.error("⚠️ Please provide SFS content to analyze.")
        elif not product_name_draft.strip():
            st.error("⚠️ Please select a product name.")
        else:
            # Mark that Find Internal Information has been completed BEFORE processing
            # This ensures the state is set immediately for the next button click
            st.session_state.find_internal_clicked = True
            
            with st.spinner("🔍 Finding internal information..."):
                try:
                    # Store in session state
                    st.session_state.product_name = product_name_draft
                    st.session_state.current_extracted_text = extracted_text
                    st.session_state.initial_analysis_done = True
                    
                    # Use the apply_prompt_file function with InternalAnalysis.md
                    internal_info = apply_prompt_file("InternalAnalysis.md", extracted_text, product_name_draft)
                    
                    # Add to conversation history
                    st.session_state.conversation_history.append({
                        "question": "Find Internal Information",
                        "answer": internal_info
                    })
                    
                    with output_container_draft:
                        st.markdown("## 🔍 Internal Information Analysis")
                        st.markdown(internal_info)

                except Exception as e:
                    st.error(f"❌ Error analyzing internal information: {str(e)}")
                    with st.expander("🐛 Error Details"):
                        st.exception(e)
    
    # Handle Generate First Draft button
    if first_draft_button:
        # Clear previous conversation context when starting First Draft workflow
        if st.session_state.conversation_history:
            # Check if last conversation was Explain SFS
            if st.session_state.conversation_history[-1].get('question') == 'Explain the SFS':
                st.session_state.conversation_history = []
        
        if not extracted_text.strip():
            st.error("⚠️ Please provide document to generate a first draft.")
        elif not product_name_draft.strip():
            st.error("⚠️ Please select a product name.")
        elif not st.session_state.find_internal_clicked:
            st.error("⚠️ Please run 'Find Internal Information' first before generating the first draft.")
        else:
            with st.spinner("✍️ Generating first draft..."):
                try:
                    # Store in session state
                    st.session_state.product_name = product_name_draft
                    st.session_state.current_extracted_text = extracted_text
                    st.session_state.initial_analysis_done = True
                    
                    # Check if there's an unanswered follow-up question/instruction in the text box
                    pending_instruction = st.session_state.get("sidebar_draft_followup_input", "").strip()
                    
                    # Build comprehensive context with internal information guidance AND all follow-up clarifications
                    context = f"Original Content:\n{extracted_text}\n\n"
                    
                    # Include the ENTIRE conversation history (internal info analysis + any follow-up clarifications)
                    if st.session_state.conversation_history or pending_instruction:
                        context += "=" * 80 + "\n"
                        context += "CONVERSATION HISTORY - Internal Information Analysis and Clarifications:\n"
                        context += "=" * 80 + "\n"
                        context += "(This includes the initial internal information identification and any follow-up clarifications)\n\n"
                        
                        # Add all exchanges in the conversation
                        for idx, exchange in enumerate(st.session_state.conversation_history, 1):
                            context += f"--- Exchange {idx} ---\n"
                            context += f"Question: {exchange['question']}\n\n"
                            context += f"Answer: {exchange['answer']}\n\n"
                        
                        # Add pending instruction/question from text box if it exists
                        if pending_instruction:
                            next_idx = len(st.session_state.conversation_history) + 1
                            context += f"--- Additional User Instruction ---\n"
                            context += f"User Directive: {pending_instruction}\n\n"
                            context += f"Note: This is an additional instruction from the user. Apply it directly when generating the first draft.\n\n"
                        
                        context += "=" * 80 + "\n\n"
                        context += "INSTRUCTIONS: The above conversation identifies internal/confidential information found in the original content, "
                        context += "along with any clarifications made during follow-up questions. "
                        context += "When generating the customer-facing first draft, you MUST:\n"
                        context += "1. Exclude or rewrite any content that relates to the internal information identified above\n"
                        context += "2. Consider ALL clarifications and follow-up discussions in the conversation\n"
                        context += "3. If there's an additional user instruction (like 'Do not include TLV'), apply it directly and strictly\n"
                        context += "4. Focus on creating documentation suitable for external customers\n"
                        context += "5. Remove implementation details, internal architecture, debug information, and any proprietary technical details\n\n"
                    
                    # Use the apply_prompt_file function with FirstDraftCTWG.md
                    first_draft = apply_prompt_file("FirstDraftCTWG.md", context, product_name_draft)
                    
                    # Add to conversation history
                    st.session_state.conversation_history.append({
                        "question": "Generate First Draft",
                        "answer": first_draft
                    })
                    
                    with output_container_draft:
                        st.markdown("## ✍️ First Draft")
                        st.markdown(first_draft)
                    
                    st.success("✅ First draft generated!")
                    
                except Exception as e:
                    st.error(f"❌ Error generating first draft: {str(e)}")
                    with st.expander("🐛 Error Details"):
                        st.exception(e)
    
    # Handle Clear button
    if clear_button_draft:
        st.session_state.conversation_history = []
        st.session_state.initial_analysis_done = False
        st.session_state.current_extracted_text = ""
        st.session_state.uploaded_file_content = ""
        # Clear follow-up answer display
        if 'last_followup_answer' in st.session_state:
            del st.session_state.last_followup_answer
        if 'last_followup_raw' in st.session_state:
            del st.session_state.last_followup_raw
        # Reset button state for Find Internal Information
        if 'find_internal_clicked' in st.session_state:
            st.session_state.find_internal_clicked = False
        # Clear text area widgets by deleting their session state keys
        if 'sidebar_extracted_text_area' in st.session_state:
            del st.session_state.sidebar_extracted_text_area
        st.rerun()
    
    # Test Section
    st.markdown("---")
    st.subheader("🧪 Capture your test results!")
    
    with st.expander("📝 Test Results", expanded=False):
        st.markdown("Capture test results for first draft generation")
        
        # Feature being tested
        first_draft_test_feature = st.text_input(
            "Feature to be tested",
            placeholder="e.g., bug analysis, bug summarize, bulk RCA, bulk bugs, First draft, resolve bug",
            help="Enter the specific feature or functionality being tested",
            key="first_draft_test_feature"
        )
        
        # Name of tester
        first_draft_tester_name = st.text_input(
            "Name of tester",
            placeholder="Enter your name",
            help="Name of the person performing the test",
            key="first_draft_tester_name"
        )
        
        # Content accuracy slider only (no location accuracy)
        first_draft_content_accuracy = st.slider(
            "Content Accuracy",
            min_value=1,
            max_value=10,
            value=10,
            help="Rate the accuracy of the content generated (1=Poor, 10=Excellent)",
            key="first_draft_test_content_accuracy"
        )
        
        # Comments text area
        first_draft_test_comments = st.text_area(
            "Comments",
            placeholder="Enter any additional comments or observations...",
            height=100,
            key="first_draft_test_comments"
        )
        
        # Wishlist text area
        first_draft_test_wishlist = st.text_area(
            "Wishlist",
            placeholder="Enter feature requests, improvements, or wishlist items...",
            height=100,
            key="first_draft_test_wishlist"
        )
        
        # Usefulness Rating
        st.markdown("---")
        first_draft_usefulness = st.radio(
            "How useful is this feature?",
            options=[
                "⛔ I'd rather do this without AI",
                "🤔 Neutral - No strong preference",
                "👍 Yes, this is useful",
                "⭐ I'd prefer CIRCUIT over manual work"
            ],
            index=2,
            help="Rate how useful you find this AI-assisted feature",
            key="first_draft_usefulness_rating"
        )
        
        # Add to Excel button
        first_draft_add_to_excel_button = st.button(
            "📊 Add to Test Excel",
            type="primary",
            use_container_width=True,
            key="first_draft_add_to_test_excel"
        )
        
        if first_draft_add_to_excel_button:
            # Import the function from sidebar_app
            from sidebar_app import save_test_results_to_excel
            import os
            
            # Get the output content from conversation history if available
            output_content = "N/A"
            if 'conversation_history' in st.session_state and st.session_state.conversation_history:
                output_content = st.session_state.conversation_history[-1]['answer']
            
            # Save to Excel with N/A for bug number and location accuracy
            try:
                save_test_results_to_excel(
                    page_name="First Draft",
                    feature=first_draft_test_feature,
                    tester_name=first_draft_tester_name,
                    bug_number="N/A",
                    output_content=output_content,
                    location_accuracy="N/A",
                    content_accuracy=first_draft_content_accuracy,
                    comments=first_draft_test_comments,
                    wishlist=first_draft_test_wishlist,
                    usefulness=first_draft_usefulness
                )
                st.success("✅ Test results saved to testresults.xlsx!")
                
                # Provide download link
                try:
                    with open("testresults.xlsx", "rb") as file:
                        st.download_button(
                            label="📥 Download testresults.xlsx",
                            data=file,
                            file_name="testresults.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key="first_draft_download_test_results"
                        )
                except:
                    pass
            except Exception as e:
                st.error(f"❌ Error saving to Excel: {str(e)}")
                with st.expander("🐛 Error Details"):
                    st.exception(e)
    
    # ===== FOLLOW-UP SECTION - Rendered after all button handlers =====
    # This ensures outputs appear before the follow-up interface
    with col2_draft:
        if st.session_state.initial_analysis_done and st.session_state.conversation_history:
            st.markdown("---")
            st.markdown("### 💬 Conversation Thread")
            st.caption("View your questions and answers, then ask follow-ups below")
            
            # Show conversation thread visibly (not in expander)
            if len(st.session_state.conversation_history) > 1:
                with st.container():
                    st.markdown("**Recent exchanges:**")
                    # Show last 3 exchanges in compact format
                    for idx, exchange in enumerate(st.session_state.conversation_history[-3:], len(st.session_state.conversation_history)-2):
                        if idx > 0:
                            st.markdown(f"**Q{idx}:** {exchange['question'][:150]}..." if len(exchange['question']) > 150 else f"**Q{idx}:** {exchange['question']}")
            
            # Compact settings
            with st.expander("⚙️ Follow-up Settings & Full History", expanded=False):
                context_window = st.slider(
                    "Context Window (number of recent exchanges to include)",
                    min_value=1,
                    max_value=10,
                    value=st.session_state.context_window_size,
                    help="Controls how many recent Q&A exchanges are included in follow-up context. Lower = faster, higher = more context.",
                    key="sidebar_draft_context_window"
                )
                st.session_state.context_window_size = context_window
                
                use_rag_followup = st.checkbox(
                    "🔍 Use RAG Search for Follow-ups",
                    value=True,
                    help="When enabled, follow-up questions will search the documentation database. When disabled, uses only conversation context.",
                    key="sidebar_draft_use_rag"
                )
                
                # Display conversation history
                st.markdown("**📜 Conversation History**")
                for idx, exchange in enumerate(st.session_state.conversation_history, 1):
                    with st.expander(f"Exchange {idx}: {exchange['question'][:80]}...", expanded=False):
                        st.markdown(f"**Question:**")
                        st.markdown(exchange['question'])
                        st.markdown(f"**Answer:**")
                        st.markdown(exchange['answer'])
            
            # Follow-up question input - more prominent
            followup_question = st.text_area(
                "Your follow-up question",
                placeholder="e.g., Can you explain the first point in more detail? Can you provide more examples?",
                height=100,
                key="sidebar_draft_followup_input"
            )
            
            ask_followup_button = st.button(
                "💬 Ask Follow-up", 
                type="primary", 
                use_container_width=True, 
                key="sidebar_draft_ask_followup"
            )
            
            # Handle follow-up button
            if ask_followup_button and followup_question.strip():
                from sidebar_app import get_relevant_context, build_followup_prompt
                from app_functions import run_agent
                
                with st.spinner("💭 Thinking..."):
                    try:
                        # Get conversation context
                        context = get_relevant_context(
                            st.session_state.conversation_history,
                            window_size=context_window
                        )
                        
                        # Build the follow-up prompt
                        full_prompt = build_followup_prompt(
                            followup_question,
                            context,
                            use_rag_followup
                        )
                        
                        # Get product name and extracted text from session state
                        product_name_state = st.session_state.get('product_name', product_name_draft)
                        extracted_text_state = st.session_state.get('current_extracted_text', extracted_text)
                        
                        # Run the agent with the follow-up prompt
                        if use_rag_followup:
                            result = run_agent(product_name_state, full_prompt, extracted_text_state)
                        else:
                            # For non-RAG, just use the LLM directly
                            from openai import OpenAI
                            client = OpenAI()
                            response = client.chat.completions.create(
                                model="gpt-4",
                                messages=[
                                    {"role": "system", "content": "You are a helpful Cisco documentation assistant."},
                                    {"role": "user", "content": full_prompt}
                                ]
                            )
                            result = {
                                'output': response.choices[0].message.content,
                                'raw': response
                            }
                        
                        # Extract answer
                        answer = result.get('output', str(result))
                        
                        # Add to conversation history
                        st.session_state.conversation_history.append({
                            "question": followup_question,
                            "answer": answer
                        })
                        
                        # Store for display
                        st.session_state.last_followup_answer = answer
                        st.session_state.last_followup_raw = result
                        
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ Error processing follow-up: {str(e)}")
                        with st.expander("🐛 Error Details"):
                            st.exception(e)
